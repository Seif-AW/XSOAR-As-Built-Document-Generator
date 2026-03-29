import yaml
import argparse
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# === Handle XSOAR non-standard YAML tags ===
class XSOARLoader(yaml.SafeLoader):
    pass

yaml.SafeLoader.add_constructor(
    'tag:yaml.org,2002:value',
    lambda loader, node: None
)

def load_playbooks(folder: Path):
    playbooks = {}
    print("\n=== Sub-Playbooks Loaded (Unique Only) ===")
    for file in folder.glob("*.yml"):
        try:
            with open(file, encoding="utf-8") as f:
                data = yaml.load(f, Loader=XSOARLoader)
                if not data or 'name' not in data:
                    continue
                name = str(data['name']).strip()
                if name not in playbooks:
                    playbooks[name] = data
                    print(f"✓ {name}")
        except:
            continue
    print(f"Total unique sub-playbooks loaded: {len(playbooks)}\n")
    return playbooks

def get_task_name(tasks, task_id):
    return tasks.get(task_id, {}).get('task', {}).get('name', f"Task {task_id}") or f"Task {task_id}"

def parse_condition(conditions):
    if not conditions:
        return ""
    parts = []
    for group in conditions:
        inner = group.get('condition', []) if isinstance(group, dict) else group
        for and_group in inner:
            if not isinstance(and_group, list):
                and_group = [and_group]
            for cond in and_group:
                if isinstance(cond, dict):
                    op = cond.get('operator', '')
                    left = cond.get('left', {}).get('value', {}).get('simple', 'value')
                    right = cond.get('right', {}).get('value', {}).get('simple', '')
                    if op == 'isNotEmpty':
                        parts.append(f"{left} is not empty")
                    elif op == 'isEqualString':
                        parts.append(f"{left} == {right}")
                    elif op == 'greaterThanOrEqual':
                        parts.append(f"{left} >= {right}")
                    elif op == 'isExists':
                        parts.append(f"{left} exists")
                    elif op:
                        parts.append(f"{left} {op} {right}")
    return " AND ".join(parts) or "Complex condition"

def format_arguments(args_dict):
    if not args_dict:
        return ""
    lines = []
    for key, val in args_dict.items():
        if isinstance(val, dict):
            if 'simple' in val:
                lines.append(f"{key}: {val['simple']}")
            elif 'complex' in val:
                root = val['complex'].get('root', '')
                transformers = val['complex'].get('transformers', [])
                trans_parts = []
                for t in transformers:
                    op = t.get('operator', '')
                    args = t.get('args', {})
                    if op == 'replace':
                        to_rep = args.get('toReplace', {}).get('value', {}).get('simple', '')
                        rep_with = args.get('replaceWith', {}).get('value', {}).get('simple', '')
                        trans_parts.append(f"replace ({to_rep} → {rep_with})")
                    elif op == 'Cut':
                        delim = args.get('delimiter', {}).get('value', {}).get('simple', '')
                        fields = args.get('fields', {}).get('value', {}).get('simple', '')
                        trans_parts.append(f"cut (delimiter:{delim} fields:{fields})")
                    elif op == 'concat':
                        prefix = args.get('prefix', {}).get('value', {}).get('simple', '')
                        suffix = args.get('suffix', {}).get('value', {}).get('simple', '')
                        parts = []
                        if prefix:
                            parts.append(f"prefix:{prefix}")
                        if suffix:
                            parts.append(f"suffix:{suffix}")
                        trans_parts.append(f"concat ({', '.join(parts)})")
                    else:
                        arg_str = ', '.join([f"{k}:{v.get('value', {}).get('simple', '')}" 
                                           for k, v in args.items() if isinstance(v, dict)])
                        trans_parts.append(f"{op} ({arg_str})" if arg_str else op)
                trans_str = " → " + " → ".join(trans_parts) if trans_parts else ""
                lines.append(f"{key}: ${{{root}}}{trans_str}")
            else:
                lines.append(f"{key}: {val}")
        else:
            lines.append(f"{key}: {val}")
    return "\n".join(lines)

def build_next_step_text(nexttasks, tasks):
    if not nexttasks:
        return "End"
    lines = []
    for label, ids in nexttasks.items():
        lbl = label.replace('#none#', 'Next').replace('#default#', 'No' if 'yes' in nexttasks else 'Default')
        for task_id in ids:
            name = get_task_name(tasks, task_id)
            lines.append(f"{lbl} → {name}")
        lines.append("")                    # ← new blank line between paths
    return "\n".join(lines).strip() or "Next"

def add_flow_table(doc, tasks, title):
    if not tasks:
        return
    doc.add_heading(title, level=2)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # === Green header (clean look) ===
    hdr = table.rows[0].cells
    header_texts = ['Condition / Task Name', 'Automation Action', 'Manual Action / Escalation', 'Next Step']

    for i, cell in enumerate(hdr):
        cell.text = header_texts[i]
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '00cc66')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run()
        run.font.name = 'Montserrat'
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(10)          # smaller = no ugly wrapping

    # === Perfect column widths (matches your clean screenshot) ===
    widths = [Pt(215), Pt(180), Pt(125), Pt(155)]
    for i, col in enumerate(table.columns):
        col.width = widths[i]

    visited = set()

    def visit(tid):
        if tid in visited: return
        visited.add(tid)
        tdata = tasks.get(tid, {})
        task = tdata.get('task', {})
        name = task.get('name', 'Untitled').strip()
        ttype = tdata.get('type', 'regular')

        if tid == "0" and not name:
            for ids in tdata.get('nexttasks', {}).values():
                for task_id in ids:
                    visit(task_id)
            return

        script = task.get('script') or task.get('scriptName') or ''
        brand = task.get('brand', '')
        args_summary = format_arguments(tdata.get('scriptarguments', {}))
        description = task.get('description', '')

        message_summary = ""
        message = tdata.get('message')
        if isinstance(message, dict):
            to = message.get('to', {}) or {}
            subject = message.get('subject', {}) or {}
            methods = ', '.join(message.get('methods', []))
            message_summary = f"To: {to.get('simple', '')}\nSubject: {subject.get('simple', '')}\nMethods: {methods}"

        next_text = build_next_step_text(tdata.get('nexttasks', {}), tasks)

        row = table.add_row().cells
        row[0].text = name

        p = row[1].paragraphs[0]
        if ttype == 'title':
            run = p.add_run("Section")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 112, 192)
        elif ttype == 'playbook':
            run = p.add_run(f"(Sub-playbook) {script}")
            run.font.color.rgb = RGBColor(255, 165, 0)
            run.bold = True
        elif ttype == 'condition':
            cond_text = parse_condition(tdata.get('conditions', []))
            run = p.add_run("(Conditional Task)")
            run.font.color.rgb = RGBColor(0, 176, 80)
            run.bold = True
            if cond_text:
                p.add_run(f"\nIf {cond_text}")
        elif ttype == 'collection':
            run = p.add_run("Data Collection Task")
            run.bold = True
            if message_summary:
                p.add_run(f"\n{message_summary}")
        elif script and ('|||' in script or task.get('iscommand')):
            cmd = script.split('|||')[-1] if '|||' in script else script
            run = p.add_run(f"{brand} → {cmd}")
            run.font.color.rgb = RGBColor(0, 112, 192)
            run.bold = True
            if args_summary:
                p.add_run(f"\n{args_summary}")
        else:
            run = p.add_run(script or name)
            run.bold = True
            if args_summary:
                p.add_run(f"\n{args_summary}")

        if ttype == 'collection':
            row[2].text = "Data Collection Task"
        elif not script and ttype == 'regular':
            row[2].text = "Manual Task"
            if description:
                row[2].text += f"\n{description[:150]}..."
        else:
            row[2].text = ""

        row[3].text = next_text

        for ids in tdata.get('nexttasks', {}).values():
            for task_id in ids:
                visit(task_id)

    visit("0")

    # === Apply Montserrat to ALL table text ===
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Montserrat'
# ====================== RUN ======================
if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("yml", type=Path)
    parser.add_argument("--subs", type=Path, default=None, help="Folder with sub-playbooks (optional)")
    parser.add_argument("--auto-subs", action="store_true", help="Auto-discover sub-playbooks in same folder as main YML")
    parser.add_argument("--out", type=Path, help="Output filename")
    parser.add_argument("--customer", default="Customer Name", help="Customer name")
    parser.add_argument("--ps", default="Your Name – Your Company", help="PS / Consultant name")
    args = parser.parse_args()
    with open(args.yml, encoding="utf-8") as f:
        main = yaml.load(f, Loader=XSOARLoader)

    if args.out:
        output_path = args.out
    else:
        output_path = args.yml.parent / f"{args.yml.stem}_As-Built.docx"

    # Auto-discover logic
    if args.auto_subs:
        subs_folder = args.yml.parent
    elif args.subs:
        subs_folder = args.subs
    else:
        subs_folder = Path(".")
    all_playbooks = load_playbooks(subs_folder)

    doc = Document()

    # === Make EVERYTHING in the document use Montserrat ===
    doc.styles['Normal'].font.name = 'Montserrat'
    for level in [1, 2, 3]:
        doc.styles[f'Heading {level}'].font.name = 'Montserrat'

    doc.add_heading('As-Built Document', 0)
    p = doc.add_paragraph(f"{args.customer} – Cortex XSOAR Automation")

    doc.add_heading('Main Playbook Flow', level=1)
    add_flow_table(doc, main.get('tasks', {}), main.get('name', 'Main Playbook'))

    # === Sub-Playbooks Details ===
    doc.add_heading('Sub-Playbooks Details', level=1)
    main_name = (main.get('name') or '').strip().lower()
    for name, sub_data in sorted(all_playbooks.items()):
        if name.strip().lower() == main_name:
            continue
        add_flow_table(doc, sub_data.get('tasks', {}), f"Sub-playbook: {name}")

    # === Document Control ===
    doc.add_heading('Document Control', level=1)
    dc = doc.add_table(rows=4, cols=2)
    from datetime import datetime
    today = datetime.now().strftime("%d-%m-%Y")
    for i, (k, v) in enumerate([
        ("Creation Date", today),
        ("Created by", f"{args.ps} - Palo Alto Networks"),
        ("Document Version", "1.0"),
        ("Reviewed and confirmed by", "")
    ]):
        dc.rows[i].cells[0].text = k
        dc.rows[i].cells[1].text = v

    # === SAVE ===
    try:
        doc.save(output_path)
        print(f"\n✅ UCD generated successfully!")
        print(f"📁 File saved to: {output_path.absolute()}")
    except Exception as e:
        print(f"\n❌ Error saving document: {e}")
